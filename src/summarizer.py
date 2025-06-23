import json
from datetime import datetime
from docx import Document
import logging

class NILNewsSummarizer:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def categorize_articles(self, articles):
        """Enhanced categorization for NIL news"""
        categories = {
            'breaking': [],      # Major announcements & breaking news
            'mega_deals': [],    # Large NIL deals (>$100k or notable)
            'policy': [],        # NCAA/State policy and regulation
            'collectives': [],   # NIL collective news
            'platforms': [],     # NIL platforms and marketplaces
            'legal': [],         # Lawsuits and legal developments
            'education': [],     # How-to guides and educational content
            'trends': []         # Market analysis and trends
        }
        
        for article in articles:
            text = (article['title'] + ' ' + article.get('summary', '')).lower()
            priority = article.get('priority_level', 'Medium')
            content_type = article.get('content_type', 'general')
            
            # Breaking news (critical priority)
            if priority == 'Critical' or any(kw in text for kw in ['breaking', 'exclusive', 'just in']):
                categories['breaking'].append(article)
            
            # Mega deals
            elif any(kw in text for kw in ['million', 'largest', 'record', 'mega deal']) and content_type == 'deals':
                categories['mega_deals'].append(article)
            
            # Policy and regulation
            elif content_type == 'policy' or any(kw in text for kw in ['ncaa rule', 'policy', 'regulation', 'compliance']):
                categories['policy'].append(article)
            
            # Collectives
            elif any(kw in text for kw in ['collective', 'booster', 'fund']):
                categories['collectives'].append(article)
            
            # Platforms and marketplaces
            elif any(kw in text for kw in ['platform', 'marketplace', 'app', 'opendorse', 'inflcr']):
                categories['platforms'].append(article)
            
            # Legal developments
            elif content_type == 'legal' or any(kw in text for kw in ['lawsuit', 'court', 'legal', 'litigation']):
                categories['legal'].append(article)
            
            # Educational content
            elif content_type == 'education' or any(kw in text for kw in ['how to', 'guide', 'tips', 'workshop']):
                categories['education'].append(article)
            
            # Market trends
            else:
                categories['trends'].append(article)
        
        return categories
    
    def create_text_summary(self, articles):
        """Create a text-based NIL news summary"""
        categories = self.categorize_articles(articles)
        
        summary = f"""# NIL Weekly News Summary
Generated on: {datetime.now().strftime('%B %d, %Y')}
Total Articles Analyzed: {len(articles)}

## Executive Summary
This week in NIL saw {len(categories['mega_deals'])} major deals, {len(categories['policy'])} policy developments, {len(categories['legal'])} legal updates, and {len(categories['collectives'])} collective/booster activities.

"""
        
        # Add key statistics if available
        deal_count = len(categories['mega_deals'])
        policy_count = len(categories['policy'])
        
        if deal_count > 0:
            summary += f"**Deal Activity:** {deal_count} significant NIL deals were announced this week.\n"
        
        if policy_count > 0:
            summary += f"**Policy Updates:** {policy_count} important policy or regulatory developments occurred.\n"
        
        summary += "\n---\n\n"
        
        # Add each category
        category_titles = {
            'breaking': '🚨 Breaking News',
            'mega_deals': '💰 Major NIL Deals',
            'policy': '📋 Policy & Regulation',
            'collectives': '🏆 Collectives & Boosters',
            'platforms': '📱 Platforms & Technology',
            'legal': '⚖️ Legal Developments',
            'education': '📚 Resources & Education',
            'trends': '📈 Market Trends & Analysis'
        }
        
        for cat_key, title in category_titles.items():
            if categories[cat_key]:
                summary += f"## {title}\n\n"
                # Show more items for breaking news and mega deals
                max_items = 10 if cat_key in ['breaking', 'mega_deals'] else 5
                
                for i, article in enumerate(categories[cat_key][:max_items], 1):
                    priority_indicator = ""
                    if article.get('priority_level') == 'Critical':
                        priority_indicator = " 🔥"
                    elif article.get('priority_level') == 'High':
                        priority_indicator = " ⭐"
                    
                    summary += f"{i}. **{article['title']}{priority_indicator}**\n"
                    summary += f"   Source: {article['source']}\n"
                    summary += f"   Date: {datetime.fromisoformat(article['published']).strftime('%B %d, %Y')}\n"
                    summary += f"   Link: {article['link']}\n"
                    
                    if article['summary']:
                        # Clean and truncate summary
                        clean_summary = article['summary'].strip()
                        if len(clean_summary) > 200:
                            clean_summary = clean_summary[:197] + '...'
                        summary += f"   Summary: {clean_summary}\n"
                    
                    if article.get('relevance_score'):
                        summary += f"   Relevance Score: {article['relevance_score']}\n"
                    
                    summary += "\n"
                
                summary += "---\n\n"
        
        # Add weekly insights
        summary += self._generate_weekly_insights(categories, articles)
        
        return summary
    
    def _generate_weekly_insights(self, categories, articles):
        """Generate insights section for the summary"""
        insights = "## Weekly Insights\n\n"
        
        # Top schools/conferences mentioned
        schools_mentioned = {}
        conferences = ['SEC', 'Big Ten', 'ACC', 'Big 12', 'Pac-12']
        
        for article in articles:
            text = (article['title'] + ' ' + article.get('summary', '')).lower()
            for conf in conferences:
                if conf.lower() in text:
                    schools_mentioned[conf] = schools_mentioned.get(conf, 0) + 1
        
        if schools_mentioned:
            insights += "**Most Active Conferences:**\n"
            for conf, count in sorted(schools_mentioned.items(), key=lambda x: x[1], reverse=True)[:3]:
                insights += f"- {conf}: {count} mentions\n"
            insights += "\n"
        
        # Key trends
        if len(categories['mega_deals']) > 2:
            insights += "**Trend Alert:** Significant increase in high-value NIL deals this week.\n\n"
        
        if len(categories['policy']) > 2:
            insights += "**Policy Watch:** Multiple policy developments suggest evolving regulatory landscape.\n\n"
        
        # Market indicators
        total_critical = sum(1 for article in articles if article.get('priority_level') == 'Critical')
        if total_critical > 3:
            insights += f"**Market Activity:** {total_critical} critical developments indicate a highly active week in the NIL space.\n\n"
        
        return insights
    
    def create_word_document(self, articles, output_path=None, custom_name=None):
        """Create a Word document NIL summary"""
        if output_path is None:
            if custom_name:
                output_path = f"outputs/{custom_name}_{datetime.now().strftime('%Y%m%d')}.docx"
            else:
                output_path = f"outputs/nil_weekly_report_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading('NIL Weekly News Summary', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Total Articles Analyzed: {len(articles)}")
        
        categories = self.categorize_articles(articles)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        exec_summary = f"This week in NIL saw {len(categories['mega_deals'])} major deals, {len(categories['policy'])} policy developments, {len(categories['legal'])} legal updates, and {len(categories['collectives'])} collective/booster activities."
        doc.add_paragraph(exec_summary)
        
        # Key highlights
        doc.add_heading('Key Highlights', level=2)
        if categories['breaking']:
            p = doc.add_paragraph()
            p.add_run('• Breaking News: ').bold = True
            p.add_run(f"{len(categories['breaking'])} critical developments this week")
        
        if categories['mega_deals']:
            p = doc.add_paragraph()
            p.add_run('• Major Deals: ').bold = True
            p.add_run(f"{len(categories['mega_deals'])} significant NIL agreements announced")
        
        # Add each category
        category_info = [
            ('breaking', '🚨 Breaking News'),
            ('mega_deals', '💰 Major NIL Deals'),
            ('policy', '📋 Policy & Regulation'),
            ('collectives', '🏆 Collectives & Boosters'),
            ('platforms', '📱 Platforms & Technology'),
            ('legal', '⚖️ Legal Developments'),
            ('education', '📚 Resources & Education'),
            ('trends', '📈 Market Trends')
        ]
        
        for cat_key, title in category_info:
            if categories[cat_key]:
                doc.add_page_break()
                doc.add_heading(title, level=1)
                
                # Add category overview
                doc.add_paragraph(f"Total articles in this category: {len(categories[cat_key])}")
                doc.add_paragraph()
                
                for i, article in enumerate(categories[cat_key][:8], 1):
                    # Article title
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. {article['title']}").bold = True
                    
                    # Metadata
                    doc.add_paragraph(f"Source: {article['source']}")
                    doc.add_paragraph(f"Date: {datetime.fromisoformat(article['published']).strftime('%B %d, %Y')}")
                    doc.add_paragraph(f"Link: {article['link']}")
                    
                    # Priority level
                    if article.get('priority_level') == 'Critical':
                        p = doc.add_paragraph()
                        p.add_run("Priority: CRITICAL").bold = True
                    
                    # Summary
                    if article['summary']:
                        clean_summary = article['summary'].strip()
                        if len(clean_summary) > 300:
                            clean_summary = clean_summary[:297] + '...'
                        doc.add_paragraph(f"Summary: {clean_summary}")
                    
                    doc.add_paragraph()  # Empty line
        
        # Add insights section
        doc.add_page_break()
        doc.add_heading('Weekly Insights & Analysis', level=1)
        insights_text = self._generate_weekly_insights(categories, articles)
        # Remove markdown formatting for Word doc
        insights_text = insights_text.replace('##', '').replace('**', '').replace('*', '•')
        doc.add_paragraph(insights_text)
        
        doc.save(output_path)
        self.logger.info(f"Word document saved to {output_path}")
        return output_path

# Usage example
if __name__ == "__main__":
    # Load articles from JSON
    with open('data/nil_articles_20250604.json', 'r') as f:
        articles = json.load(f)
    
    summarizer = NILNewsSummarizer()
    
    # Create text summary
    text_summary = summarizer.create_text_summary(articles)
    with open('outputs/nil_summary.txt', 'w') as f:
        f.write(text_summary)
    
    # Create Word document
    summarizer.create_word_document(articles)